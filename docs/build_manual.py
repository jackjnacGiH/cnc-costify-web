"""
Build CNC Costify AI V5.0 user manual (Thai) as a .docx with embedded screenshots.
Uses python-docx. Screenshots are at docs/manual_images/01..09.png
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(DOCS_DIR, 'manual_images')
OUT = os.path.join(DOCS_DIR, 'CNC_Costify_AI_V5.0_User_Manual.docx')

doc = Document()

# Default font: TH Sarabun New for Thai readability, fall back to Tahoma
def _set_default_font():
    style = doc.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'TH Sarabun New')
    rfonts.set(qn('w:hAnsi'), 'TH Sarabun New')
    rfonts.set(qn('w:cs'), 'TH Sarabun New')
    rfonts.set(qn('w:eastAsia'), 'TH Sarabun New')

_set_default_font()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def H1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return p

def H2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    return p

def H3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    return p

def P(text, bold=False, size=14, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(14)
    return p

def NUM(text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(14)
    return p

def IMG(filename, caption=None, width_inches=6.5):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        P(f'[ภาพ {filename} หาไม่พบ]', color=(0xC0, 0x39, 0x2B))
        return
    doc.add_picture(path, width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(13)
        r.italic = True
        r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

def NOTE(text, kind='info'):
    """Coloured callout box (info/warn/tip)."""
    colors = {
        'info': (0xE0, 0xF2, 0xFE, 0x07, 0x4D, 0xC8),  # bg-blue, text-blue
        'warn': (0xFE, 0xF3, 0xC7, 0xB4, 0x53, 0x09),
        'tip':  (0xDC, 0xFC, 0xE7, 0x05, 0x96, 0x69),
        'err':  (0xFE, 0xE2, 0xE2, 0xC0, 0x39, 0x2B),
    }
    bg_r, bg_g, bg_b, fg_r, fg_g, fg_b = colors.get(kind, colors['info'])
    icon = {'info': 'ℹ️', 'warn': '⚠️', 'tip': '💡', 'err': '❌'}.get(kind, '')
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # Shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f'{bg_r:02X}{bg_g:02X}{bg_b:02X}')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    r = p.add_run(f'{icon} {text}')
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(fg_r, fg_g, fg_b)
    r.bold = True
    return table

def PAGEBREAK():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ────────────────────────────────────────────────────────────────────────
# Cover page
# ────────────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('\n\n\n\n')
r.font.size = Pt(20)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('CNC Costify AI')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(48)
r.bold = True
r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Version 5.0')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('คู่มือการใช้งาน (User Manual)')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph('\n' * 3)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('โปรแกรมคำนวณราคางาน CNC ด้วย AI\nรองรับไฟล์ STEP / PDF / JPG พร้อมระบบจัดเก็บข้อมูล Excel')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(20)
r.italic = True
r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

doc.add_paragraph('\n' * 6)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(f'จัดทำเมื่อ: {datetime.now().strftime("%d/%m/%Y")}\nwww.cnccostify.cloud')
r.font.name = 'TH Sarabun New'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# Table of Contents
# ────────────────────────────────────────────────────────────────────────
H1('สารบัญ')

toc_items = [
    ('บทที่ 1', 'การติดตั้งโปรแกรม'),
    ('บทที่ 2', 'หน้าหลัก — คำนวณจากไฟล์ STEP'),
    ('บทที่ 3', 'PDF/JPG Costify — คำนวณหลายไฟล์ด้วย AI'),
    ('บทที่ 4', 'หน้าตั้งค่า'),
    ('บทที่ 5', 'การจัดการสิทธิ์การใช้งาน'),
    ('บทที่ 6', 'แชทปรึกษา AI'),
    ('บทที่ 7', 'การ Export ข้อมูลเป็น Excel'),
    ('บทที่ 8', 'แก้ปัญหาเบื้องต้น (Troubleshooting)'),
]
for ch, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f'{ch}  ')
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(15)
    r.bold = True
    r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    r2 = p.add_run(title)
    r2.font.name = 'TH Sarabun New'
    r2.font.size = Pt(15)

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 1 — การติดตั้งโปรแกรม
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 1 — การติดตั้งโปรแกรม')
P('CNC Costify AI V5.0 มาในรูปแบบ One-Click Installer ขนาดประมาณ 660 MB ซึ่งบรรจุ Library, Plugin และ Parameter ที่จำเป็นทั้งหมดไว้ในไฟล์เดียว ไม่ต้องติดตั้ง Python, Node.js หรือ dependencies อื่นๆ เพิ่มเติม')

H2('1.1 ความต้องการของระบบ (System Requirements)')
BULLET('ระบบปฏิบัติการ: Windows 10 หรือ Windows 11 (64-bit)')
BULLET('RAM: ขั้นต่ำ 4 GB (แนะนำ 8 GB ขึ้นไป)')
BULLET('พื้นที่ว่าง Hard Disk: อย่างน้อย 2 GB')
BULLET('การเชื่อมต่ออินเทอร์เน็ต: จำเป็นสำหรับการเรียก AI (Gemini / OpenRouter) และดึงอัตราแลกเปลี่ยน USD→THB')

H2('1.2 ขั้นตอนการติดตั้ง')
NUM('คลิกขวาที่ไฟล์ "CNC Costify AI V5.0 Setup.exe" → เลือก "Run as administrator"')
NUM('เมื่อ Windows แสดง UAC prompt ให้กด "Yes" เพื่ออนุญาตการติดตั้ง')
NUM('Installer จะติดตั้งแบบอัตโนมัติ (One-Click) ใช้เวลาประมาณ 1-2 นาที')
NUM('เมื่อเสร็จสิ้น ระบบจะเปิดโปรแกรมให้อัตโนมัติพร้อมสร้าง shortcut บน Desktop และ Start Menu')

NOTE('หากติดตั้ง V5.0 รุ่นเก่าอยู่แล้ว Installer จะหยุดโปรเซสที่กำลังรันและแทนที่ไฟล์ให้อัตโนมัติ ไม่ต้องถอนการติดตั้งเก่าก่อน', 'tip')

H2('1.3 การเริ่มต้นใช้งานครั้งแรก')
NUM('เปิดโปรแกรมจาก Desktop shortcut "CNC Costify AI V5.0"')
NUM('รอจนสถานะมุมขวาบนแสดง "พร้อมใช้งาน" (จุดเขียว)')
NUM('สลับภาษาได้ที่ปุ่ม TH/EN มุมขวาบน')
NUM('ก่อนใช้งานครั้งแรก ต้องตั้งค่า API Key ก่อน — ดูบทที่ 4')

NOTE('Backend Python (พร้อม OCC libraries) จะถูกเปิดอัตโนมัติบน port 5001 — ไม่ต้องกดเปิดเอง', 'info')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 2 — หน้าหลัก
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 2 — หน้าหลัก (คำนวณจากไฟล์ STEP)')
P('หน้าหลักใช้สำหรับคำนวณราคาชิ้นงาน CNC จากไฟล์ STEP โดยระบบจะอ่านปริมาตรชิ้นงานและขนาด Stock Material อัตโนมัติด้วย OpenCASCADE (OCC)')

IMG('01_home_empty.png', 'ภาพที่ 2.1 — หน้าหลักในสถานะเริ่มต้น')

H2('2.1 ส่วนประกอบของหน้าหลัก')
BULLET('แถบเมนูบน: หน้าหลัก / PDF/JPG / ตั้งค่า / สิทธิ์การใช้งาน / แชท AI')
BULLET('สลับภาษา TH/EN และไฟแสดงสถานะ Backend (พร้อมใช้งาน/Online)')
BULLET('ฝั่งซ้าย: ฟอร์มอินพุต (อัปโหลด STEP, Material, Coating, Surface, Wire Cut, Quantity)')
BULLET('ฝั่งขวา: Panel "ผลลัพธ์การคำนวณ" แสดงราคาวัสดุ, CNC, Coating, Setup, ราคา/ชิ้น, ราคารวม')
BULLET('ปุ่มล่าง: บันทึกข้อมูล / คัดลอก / เริ่มใหม่')

H2('2.2 ขั้นตอนการคำนวณ')
NUM('คลิก "อัปโหลด" และเลือกไฟล์ .step (.stp ก็ได้)')
NUM('รอประมาณ 1-2 วินาที — ระบบจะอ่าน Volume, Stock Size, Stock Volume, Shape อัตโนมัติ')
NUM('กรอก Material (เช่น A6061, S45C) — ระบบมีฐานข้อมูลวัสดุ 100+ รายการ')
NUM('เลือก Coating (Anodize, Hardening, etc.) ถ้ามี')
NUM('เลือก Surface (พ่นทราย etc.) และ Wire Cut ถ้ามี')
NUM('ระบุ Precision Base และจำนวน (Quantity)')
NUM('กดปุ่ม "คำนวณราคา" สีน้ำเงิน — ผลลัพธ์จะปรากฏทางขวา')

IMG('02_home_with_results.png', 'ภาพที่ 2.2 — หน้าหลักหลังคำนวณราคา')

H2('2.3 การอ่านผลลัพธ์')
P('ตัวอย่างจากภาพที่ 2.2:', bold=True)
BULLET('ปริมาตรชิ้นงาน: 4,368,222.94 mm³ (อ่านจาก STEP โดย OCC)')
BULLET('Stock Size: 25.00 × 300.00 × 920.00 mm (Bounding Box)')
BULLET('Material A6061 → Material Cost: 4,660 บาท (รวม starting price)')
BULLET('Coating Anodize → 950.15 in² × 3,330 บาท')
BULLET('CNC Process: 7,600 บาท + Setup 350 บาท')
BULLET('ราคาชิ้นงาน: 15,765 บาท/Ea. × 2 ชิ้น = ราคารวม 31,530 บาท')

NOTE('ค่า "Round up" หมายถึงปัดขึ้นเลขทศนิยมเศษหรือปัดสิบบาท ระบบจะใช้สูตรของบริษัทอัตโนมัติ', 'info')

H2('2.4 ปุ่มจัดการ')
BULLET('บันทึกข้อมูล (ปุ่มน้ำเงิน): บันทึกผลลัพธ์ลงไฟล์ Excel ที่ตั้งค่าไว้')
BULLET('คัดลอก (ปุ่มเขียว): คัดลอกราคาเป็น text สำหรับวางใน email/quotation')
BULLET('เริ่มใหม่ (ปุ่มแดง): เคลียร์ฟอร์มทั้งหมด')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 3 — PDF/JPG Costify
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 3 — PDF/JPG Costify (คำนวณหลายไฟล์ด้วย AI)')
P('หน้านี้รองรับการประมวลผลแบบ Batch สำหรับไฟล์ PDF หรือ JPG โดยใช้ AI วิเคราะห์ข้อมูลในแบบ (drawing) และคำนวณราคาให้พร้อมกันทุกไฟล์')

IMG('03_pdfjpg_empty.png', 'ภาพที่ 3.1 — หน้า PDF/JPG Costify ในสถานะว่าง')

H2('3.1 ส่วนประกอบ')
BULLET('คิวไฟล์ (ซ้าย): drag & drop หรือคลิก "เพิ่มไฟล์แบบงาน" — รองรับ PDF / JPG / PNG')
BULLET('ปุ่ม "ประมวลผล X รายการ" สำหรับเริ่มวิเคราะห์ทุกไฟล์ในคิว')
BULLET('ปุ่ม "เคลียร์ข้อมูลทั้งหมด" สำหรับลบไฟล์และผลลัพธ์ทั้งหมด')
BULLET('Precision Base × ค่าตัวคูณ (×1.0 = มาตรฐาน)')
BULLET('สถานะวัสดุ: 108 รายการ (โหลดแล้ว)')
BULLET('ปุ่ม EXPORT EXCEL ขวาบน — บันทึกผลทั้งหมดลง Excel')

H2('3.2 ขั้นตอนการใช้งาน')
NUM('ลากไฟล์ PDF/JPG เข้าพื้นที่คิวไฟล์ (รองรับหลายไฟล์พร้อมกัน)')
NUM('ตรวจรายการในคิว — สามารถลบไฟล์เดี่ยวด้วยไอคอนถังขยะ')
NUM('กดปุ่ม "ประมวลผล X รายการ"')
NUM('รอ AI วิเคราะห์ — แต่ละไฟล์จะเปลี่ยนสถานะเป็น "สำเร็จ" + ผลในตารางสรุปจะขึ้น')
NUM('ตรวจ/แก้ไขข้อมูล: Material, Size Mat, Coating, จำนวน')
NUM('Export เป็น Excel เมื่อเสร็จ')

IMG('04_pdfjpg_analyzed.png', 'ภาพที่ 3.2 — ผลลัพธ์การวิเคราะห์ AI พร้อม Engine + Cost badge')

H2('3.3 Badge ที่แสดงในแต่ละแถว')
BULLET('Engine badge (สีฟ้า/ส้ม): บอกว่า AI ตัวไหนวิเคราะห์ — ☁️ Gemini หรือ 🔀 OpenRouter (model name)')
BULLET('Cost badge (สีเขียว 💰 X.XX บาท): ค่าใช้จ่าย Token ของไฟล์นั้น แปลงจาก USD → THB อัตโนมัติ (อัตรา live จาก frankfurter.app)')
BULLET('AI ประมาณ (ป้ายเหลือง): ขนาด Stock Mat ที่ AI ทำนายจากแบบ — แก้ไขได้')
BULLET('ราคาขั้นต่ำ (สีส้ม): กรณีค่าวัสดุต่ำกว่า starting price ที่ตั้งไว้ ระบบจะใช้ starting price แทน')
BULLET('บันทึกเทคนิค: Note จาก AI เกี่ยวกับเทคนิคการผลิต/รูเจาะพิเศษ ฯลฯ')

NOTE('Cost ของแต่ละไฟล์คำนวณจาก: (input_tokens × ราคา_in + output_tokens × ราคา_out) × อัตรา USD→THB ราคา per token ดูจาก openrouter.ai/models และ ai.google.dev/pricing', 'info')

H2('3.4 ฐานข้อมูลโมเดล AI ที่รองรับ')
BULLET('Qwen 2.5-VL 72B Instruct (paid, ~0.40$/1M in)')
BULLET('GLM-4.5V (Chinese — z-ai/glm-4.5v) (paid)')
BULLET('Claude 3.5 Sonnet (anthropic) (paid)')
BULLET('Gemini 2.0 Flash Exp (free)')
BULLET('Llama 3.2 90B Vision (free)')
BULLET('Gemini 2.5 Flash / 1.5 Flash / 2.0 Flash (จาก Google API ตรง)')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 4 — หน้าตั้งค่า
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 4 — หน้าตั้งค่า')
P('หน้าตั้งค่าใช้กำหนดค่าทุกอย่างที่เกี่ยวข้องกับการคำนวณ — ตั้งแต่ราคา Process, ฐานข้อมูลวัสดุ, Coating, Surface, Wire Cut, ราคาเริ่มต้น (starting price), Cutting Group, Excel path และ API ของ AI')

IMG('05_settings_main.png', 'ภาพที่ 4.1 — หน้าตั้งค่าแบบ Compact 3 คอลัมน์')

H2('4.1 คอลัมน์ที่ 1 — Process Pricing & ระบบ')
BULLET('Process Pricing: ราคา/หน่วย ของกระบวนการ CNC ต่างๆ (กัดเข้า, บานหยาบ, ค่ามุดเจาะ, กัดเฉพาะ, กัดงานมาก)')
BULLET('ตั้งค่าไฟล์บันทึก Excel: path ที่จะใช้บันทึกผลลัพธ์ (เช่น D:\\Costify Data.xlsx)')
BULLET('จัดการ Cutting Group: เพิ่ม/แก้ไข/ลบกลุ่มการตัด')
BULLET('ตั้งค่า API (การ์ดสีม่วง): เปิด Modal ตั้งค่า API — ดู 4.4')

H2('4.2 คอลัมน์ที่ 2 — Material & Coating')
BULLET('จัดการข้อมูลและราคาวัสดุ: เพิ่มชื่อวัสดุ, Density, ราคา (บาท/Kg หรือ บาท/mm³), กลุ่มการตัด')
BULLET('ตั้งราคา Coating: เลือกประเภท Coating (Anodize, Hardening, Black Oxide ฯลฯ) + ราคา')

H2('4.3 คอลัมน์ที่ 3 — Surface, Starting Price & Backup')
BULLET('Surface: ตั้งราคา Surface treatment เช่น พ่นทราย (บาท/in²)')
BULLET('Wire Cut: ราคาต่อ mm')
BULLET('Starting Price: ราคาขั้นต่ำสำหรับ Material Cost / Coating / CNC+Process / CNC Setup')
BULLET('Backup/Restore: สำรองและคืนค่าตั้งค่าทั้งหมด')

H2('4.4 ตั้งค่า API (สำคัญที่สุด)')
P('คลิกปุ่ม "ตั้งค่า API" สีม่วงเพื่อเปิด Modal:')

IMG('06_api_settings_modal.png', 'ภาพที่ 4.2 — Modal ตั้งค่า API Provider')

H3('Gemini API Keys (Provider หลัก)')
BULLET('ใส่ API Key ของ Google Gemini ได้หลายตัว — บรรทัดละ 1 key')
BULLET('ระบบจะเวียน key อัตโนมัติเมื่อ key หนึ่ง quota เต็ม (สถานะ "key 1/3" บนแถบบน)')
BULLET('สมัครฟรีที่ aistudio.google.com — 60 req/min ต่อ key')
BULLET('Gemini Model: เลือกโมเดล (gemini-2.5-flash แนะนำ)')
BULLET('ปุ่ม "โหลดรายชื่อ" — ดึงรายการโมเดลที่ใช้ได้จาก API ตรง')

H3('OpenRouter Fallback (สำรองอัตโนมัติ)')
BULLET('ใช้เมื่อ Gemini ทุก key quota เต็ม')
BULLET('สมัครฟรีที่ openrouter.ai — รับเครดิตทดลองใช้')
BULLET('Model: เลือกได้หลายตัว (Qwen 2.5-VL, GLM-4.5V, Claude 3.5 Sonnet, Llama 3.2 90B Vision)')
BULLET('ทำงานเป็น Fallback อัตโนมัติเมื่อ Gemini ใช้ไม่ได้ — แอปแสดง badge 🔀 OpenRouter')

NOTE('โปรแกรมจะใช้โมเดลตามที่เลือกใน Settings — ไม่มีการ silent retry ด้วยโมเดลอื่นโดยอัตโนมัติ ถ้า model ที่เลือกล้มเหลว แอปจะแสดง error ให้เห็นชัดเจน', 'tip')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 5 — สิทธิ์การใช้งาน
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 5 — การจัดการสิทธิ์การใช้งาน')
P('โปรแกรมใช้ระบบ License แบบผูกกับ Hardware ID ของเครื่อง — ป้องกันการคัดลอกโดยไม่ได้รับอนุญาต')

IMG('07_license.png', 'ภาพที่ 5.1 — หน้าจัดการสิทธิ์การใช้งาน')

H2('5.1 ข้อมูลสิทธิ์')
BULLET('สถานะสิทธิ์: ใช้งานได้ (เขียว) / หมดอายุ / ไม่พบไฟล์สิทธิ์')
BULLET('วันคงเหลือ: จำนวนวันก่อนหมดอายุ')
BULLET('ใช้ได้ถึง: วันสิ้นสุดสิทธิ์')
BULLET('Hardware ID: รหัสที่สร้างจาก hardware ของเครื่อง — ใช้สำหรับขอ License Key (ปุ่มคัดลอก)')

H2('5.2 การขอ License Key')
NUM('คัดลอก Hardware ID จากหน้าสิทธิ์การใช้งาน')
NUM('ติดต่อผู้ดูแลผ่านช่องทางที่แสดง:')
BULLET('โทรศัพท์: 08 1144 2000')
BULLET('อีเมล: info@cnccostify.cloud')
BULLET('เว็บไซต์: www.cnccostify.cloud')
NUM('ส่ง Hardware ID + ชื่อผู้ติดต่อ/บริษัท')
NUM('ผู้ดูแลจะส่งไฟล์สิทธิ์ (.json หรือ .dat) กลับมาทาง email')

H2('5.3 การติดตั้งสิทธิ์')
NUM('คลิก "เลือกไฟล์" → เลือกไฟล์สิทธิ์ที่ได้รับ')
NUM('เลือก "ตำแหน่งติดตั้ง" — แนะนำ "โฟลเดอร์ผู้ใช้ (Roaming)" สำหรับการใช้งานปกติ')
NUM('คลิก "นำเข้าสิทธิ์การใช้งาน"')
NUM('คลิก "ตรวจสอบสถานะ" เพื่อยืนยัน — ควรเปลี่ยนเป็น "ใช้งานได้" สีเขียว')

NOTE('สิทธิ์การใช้งานผูกกับ Hardware ID — ถ้าเปลี่ยนคอมพิวเตอร์ต้องขอสิทธิ์ใหม่', 'warn')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 6 — แชทปรึกษา AI
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 6 — แชทปรึกษา AI')
P('ฟีเจอร์ AI Chat ใช้สำหรับสอบถามเรื่องวัสดุ, สเปค, การทดแทนวัสดุ, ความหนาแน่น (Density), Coating, Heat Treatment, Machining Parameters และคำปรึกษาทั่วไปด้านวิศวกรรมการผลิต — ระบบเลือก AI Provider อัตโนมัติ (Gemini ก่อน → OpenRouter Fallback) และ AI ถูกปรับให้ตอบสั้น+กระชับ+เร็ว เพื่อให้ใช้งานในไลน์การผลิตได้ทันท่วงที')

IMG('08_ai_chat.png', 'ภาพที่ 6.1 — หน้าแชทปรึกษา AI พร้อมตัวอย่างคำตอบ Speed/Feed/DOC')

H2('6.1 ส่วนประกอบ')
BULLET('หัวข้อคำถาม (ตัวอย่าง): Dropdown 13 หัวข้อพร้อม emoji icon — ช่วยให้ user เริ่มได้ทันที')
BULLET('คำแนะนำ: tip การถาม ให้ระบุบริบท (วัสดุ, กระบวนการ, ความแม่นยำ, สภาพการใช้งาน)')
BULLET('พื้นที่ chat: แสดงประวัติการสนทนา (สีฟ้า=ผู้ใช้, สีเทา=AI)')
BULLET('Input box: พิมพ์คำถาม + ปุ่ม ส่งข้อความ / เริ่มใหม่ / ล้างข้อความ')
BULLET('Density Auto-Detect: เมื่อเลือกหัวข้อ Density ระบบจะดึงค่าวัสดุที่กำหนดในหน้าหลักมาแสดงอัตโนมัติ')

NOTE('ไม่ต้องเลือก AI Provider เอง — ระบบ route Gemini → OpenRouter Fallback อัตโนมัติเมื่อ Gemini quota เต็ม', 'tip')

H2('6.2 หัวข้อคำถาม 13 ประเภท')
P('แต่ละหัวข้อมี Default Prompt ตัวอย่าง user สามารถแก้ไขเพิ่มเติมก่อนกดส่งได้:')

IMG('08b_chat_dropdown.png', 'ภาพที่ 6.2 — Dropdown แสดงหัวข้อคำถามทั้ง 13 ประเภท', width_inches=5.5)

BULLET('💬 ปรึกษาทั่วไป (default) — สำหรับคำถามทั่วไปที่ไม่อยู่ในหมวด')
BULLET('⚖️ ความหนาแน่น (Density) — ตอบ g/cm³ + ตัวอย่างใช้งาน')
BULLET('📋 สเปควัตถุดิบ / มาตรฐานเทียบ — ส่วนผสม, ASTM/JIS/DIN, Hardness')
BULLET('🔁 วัสดุทดแทน — แนะนำ 2-3 ตัว + เหตุผลสั้นๆ')
BULLET('🔥 การชุบแข็ง / Heat Treatment — วิธี, อุณหภูมิ, HRC, ข้อควรระวัง')
BULLET('✨ Surface Treatment / ผิวงาน — ตัวเลือก, ความหนา, ราคาประมาณการ')
BULLET('🎨 เลือก Coating ให้เหมาะกับงาน — Anodize / Hard / Black Oxide ฯลฯ')
BULLET('⚙️ พารามิเตอร์ตัดเฉือน (Speed/Feed) — สำหรับ Endmill / Drill')
BULLET('🌡️ ทนความร้อน — อุณหภูมิสูงสุด + วัสดุเทียบเคียง')
BULLET('🧪 ทนการกัดกร่อน — เกรด/coating ที่ทน + สภาพแวดล้อม')
BULLET('💪 คุณสมบัติทางกล (Tensile/Yield) — Tensile, Yield, Elongation, Hardness')
BULLET('⚡ การเชื่อม / Welding — วิธีที่เหมาะ, ลวดเชื่อม, Pre/post heat')
BULLET('📐 Tolerance / ความเที่ยง — ตามเกรด ISO 2768 หรือ Class fit')

H2('6.3 ตัวอย่างการใช้งานจริง')
P('ผู้ใช้ถาม: "พารามิเตอร์ตัด S45C ด้วย Endmill HSS Ø10 — Speed/Feed/DOC"', bold=True)
P('AI ตอบ (ตัวอย่างจากภาพที่ 6.1):')
BULLET('Speed: 30–40 m/min')
BULLET('Feed: 0.1–0.15 mm/tooth')
BULLET('DOC (Radial): ≤5 mm (50% Ø)')
BULLET('DOC (Axial): 2–3 mm')

NOTE('AI ตอบสั้น กระชับ ≤6 bullets / ≤100 คำ — ใช้ thinkingBudget=0 + maxOutputTokens=1024 เพื่อความเร็วสูงสุด (~2 วินาที/ครั้ง)', 'info')

H2('6.4 ตัวอย่างคำถามอื่นๆ ที่ใช้บ่อย')
BULLET('"การชุบแข็ง S45C ให้ HRc 50-55 ใช้วิธีอะไร"')
BULLET('"วัสดุทดแทน S45C ที่ราคาถูกกว่า แต่คุณสมบัติใกล้เคียง"')
BULLET('"งานทนความร้อน 500°C ควรใช้วัสดุอะไร"')
BULLET('"เลือก Coating ระหว่าง Anodize Hard กับ Black Oxide สำหรับงานทนการสึกหรอ"')
BULLET('"Tolerance ตาม ISO 2768-mK สำหรับขนาด 50 mm คือเท่าไหร่"')
BULLET('"การเชื่อม A6061-T6 — วิธีไหนเหมาะ ลวดเชื่อมเกรดใด ต้อง Pre-heat ไหม"')

NOTE('AI ใช้ฐานความรู้ Materials Engineering จาก ai_skills/ ที่บันเดิ้ลในแอป — ตอบเป็นภาษาไทยได้', 'info')

H2('6.5 ข้อควรระวัง')
BULLET('โปรดหลีกเลี่ยงข้อมูลอ่อนไหว (เลขบัญชี, รหัสลูกค้า)')
BULLET('คำตอบจาก AI เป็น reference เท่านั้น ควรตรวจสอบกับ datasheet จริงเสมอ')
BULLET('การแชทใช้ token จะมีค่าใช้จ่าย (กรณีใช้ paid model) — ราคา/ครั้งโดยทั่วไป < 0.05 บาท')
BULLET('ถ้า Gemini quota เต็ม ระบบสลับ OpenRouter อัตโนมัติ — ผู้ใช้ไม่ต้องทำอะไร')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 7 — Excel Export
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 7 — การ Export ข้อมูลเป็น Excel')
P('โปรแกรมรองรับการบันทึกผลลัพธ์ลงไฟล์ Excel ทั้งจากหน้าหลัก (รายชิ้น) และ PDF/JPG (batch ทีละหลายชิ้น)')

H2('7.1 การตั้งค่า Path Excel')
NUM('ไปที่ ตั้งค่า → ตั้งค่าไฟล์บันทึก Excel')
NUM('คลิก "เลือกไฟล์..." แล้วเลือกตำแหน่งและชื่อไฟล์ (เช่น D:\\Costify Data.xlsx)')
NUM('คลิก "บันทึกแผ่นไฟล์" เพื่อยืนยัน path')

H2('7.2 การ Export จากหน้าหลัก')
NUM('คำนวณราคาเสร็จ → คลิกปุ่ม "บันทึกข้อมูล" สีน้ำเงิน')
NUM('ระบบจะเพิ่ม sheet ใหม่ในไฟล์ Excel ที่ตั้งไว้')
NUM('Sheet name format: "Main DD-MM-YY HH.MM"')

H2('7.3 การ Export จาก PDF/JPG')
NUM('วิเคราะห์ไฟล์เสร็จ → คลิก "EXPORT EXCEL" (ขวาบน)')
NUM('ระบบสร้าง sheet ใหม่ format: "PDF DD-MM-YY HH.MM"')
NUM('Format ของ sheet:')
BULLET('Column A: No. (เลขลำดับ)')
BULLET('Column B: ชื่องาน / Part Name')
BULLET('Column C-N: Part No, Drawing No, วัสดุ, ขนาด, จำนวน, ราคาแต่ละ process')
BULLET('Column O: หมายเหตุ (Wrap text — รองรับข้อความยาว)')
BULLET('แถวสุดท้าย: ยอดรวมทั้งหมด (เซลล์ไฮไลต์เหลือง + double underline)')

IMG('09_excel_export.png', 'ภาพที่ 7.1 — Excel ที่ Export จาก PDF/JPG Costify')

H2('7.4 รูปแบบการ Format ตัวเลข')
BULLET('ตัวเลขทุก column: format #,##0.00 (เช่น 1,290.00)')
BULLET('Header row: bold สีขาวบนพื้นน้ำเงินเข้ม + center')
BULLET('No. column: center-aligned')
BULLET('Notes column: wrap text กว้าง 50 ตัวอักษร')
BULLET('Total cell: bold, ชิดขวา, พื้นเหลือง #FFFF00, double-underline border')

NOTE('ถ้าไฟล์ Excel เปิดอยู่ใน Microsoft Excel ขณะ Export จะได้ error "FILE_BUSY" — ปิดไฟล์ก่อนแล้วลองใหม่', 'warn')

PAGEBREAK()

# ────────────────────────────────────────────────────────────────────────
# บทที่ 8 — Troubleshooting
# ────────────────────────────────────────────────────────────────────────
H1('บทที่ 8 — แก้ปัญหาเบื้องต้น (Troubleshooting)')

H2('8.1 สถานะ Backend ไม่ขึ้น "พร้อมใช้งาน"')
BULLET('สาเหตุ: Python backend ไม่ได้เริ่มทำงาน หรือ port 5001 ถูกใช้โดยโปรแกรมอื่น')
BULLET('วิธีแก้: ปิดโปรแกรมแล้วเปิดใหม่ (Backend จะ restart อัตโนมัติ)')
BULLET('ถ้ายังไม่ได้: ตรวจ Antivirus/Firewall ที่อาจ block CNC-Costify-AI.exe')

H2('8.2 อัปโหลด STEP แล้ว Volume คำนวณไม่ถูก/แสดง error')
BULLET('สาเหตุ: ไฟล์ STEP เสีย หรือ OCC backend ไม่ทำงาน')
BULLET('ผลลัพธ์: แสดง error สีแดง "❌ คำนวณ STEP ไม่สำเร็จ — pythonocc-core not installed" หรือ "OCCT backend ไม่พร้อม"')
BULLET('วิธีแก้: ปิด-เปิดโปรแกรมใหม่ → ตรวจไฟล์ STEP ด้วยโปรแกรม CAD อื่น → ติดตั้ง V5.0 ใหม่')

NOTE('โปรแกรมจะไม่ใช้ JS heuristic ถ้า OCC ใช้งานไม่ได้ — แสดง error ชัดเจนแทนการคำนวณค่าที่ผิดพลาด', 'tip')

H2('8.3 PDF/JPG วิเคราะห์ไม่ได้')
BULLET('สาเหตุ: API Key ไม่ถูกตั้ง หรือ quota เต็มหรือ model ที่เลือกไม่ support PDF')
BULLET('Error ที่อาจเห็น:')
BULLET('"API key not valid" → ตรวจ key ในตั้งค่า API')
BULLET('"OpenRouter ตอบกลับ error: Provider returned error" → ลองเปลี่ยน model')
BULLET('"Provider returned error" + กับไฟล์ PDF → ใช้ model ที่ support PDF native (Qwen 2.5-VL หรือ Claude)')

H2('8.4 Export Excel ไม่ได้ (FILE_BUSY)')
BULLET('สาเหตุ: ไฟล์ Excel เปิดอยู่ใน Microsoft Excel')
BULLET('วิธีแก้: ปิดไฟล์ใน Excel ก่อนแล้วลอง Export ใหม่')

H2('8.5 ลบ Gemini Key ไม่ออก')
BULLET('ในตั้งค่า API: ล้าง textarea Gemini API Keys ให้ว่าง → กด "บันทึก"')
BULLET('ระบบจะส่ง empty string ไป backend ซึ่งจะลบ key ทั้งหมด')
BULLET('เปิด Modal ใหม่ — textarea ควรว่างเปล่า')

H2('8.6 ข้อมูลคู่มือเพิ่มเติม')
BULLET('คู่มือ Admin: docs/Admin_User_Guide_th.html')
BULLET('License Management: docs/License_Management_User_Guide_th.md')
BULLET('Issue Tracker: github.com/anthropics/claude-code/issues (สำหรับ bug report)')

# ────────────────────────────────────────────────────────────────────────
# Footer / End
# ────────────────────────────────────────────────────────────────────────
PAGEBREAK()
H1('— จบคู่มือ —')
P('ขอบคุณที่ใช้งาน CNC Costify AI V5.0', bold=True, size=18, color=(0x25, 0x63, 0xEB))
P('หากมีคำถามหรือพบปัญหา สามารถติดต่อผู้ดูแลระบบได้ที่:')
BULLET('โทรศัพท์: 08 1144 2000')
BULLET('อีเมล: info@cnccostify.cloud')
BULLET('เว็บไซต์: www.cnccostify.cloud')

# Save
doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)/1024:.1f} KB')
